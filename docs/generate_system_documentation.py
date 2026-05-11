from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\VincentOmbok\Documents\Kura27project")
OUTPUT = ROOT / "docs" / "KURA_LIVE_System_Documentation.docx"


def set_cell_text(cell, text, bold=False, font_size=10.5, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2F3")
        borders.append(element)
    tbl_pr.append(borders)


def add_table(document, headers, rows, column_widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)

    header_row = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_row[idx].width = Inches(column_widths[idx])
        set_cell_text(header_row[idx], header, bold=True, font_size=10.5, color=RGBColor(255, 255, 255))
        shade_cell(header_row[idx], "244061")

    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].width = Inches(column_widths[idx])
            set_cell_text(row[idx], value)
    return table


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    return paragraph


def add_body(document, text, bold_lead=None):
    paragraph = document.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        lead.font.name = "Arial"
        rest = paragraph.add_run(text[len(bold_lead):])
        rest.font.name = "Arial"
    else:
        run = paragraph.add_run(text)
        run.font.name = "Arial"
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return paragraph


def add_code_block(document, lines):
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)


def set_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    title = document.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(12)

    for level, size in ((1, 16), (2, 14), (3, 12)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_styles(document)

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_paragraph.add_run("KURA LIVE • System Reference Guide")
    header_run.font.name = "Arial"
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(91, 155, 213)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_paragraph.add_run("KURA LIVE Documentation")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("KURA LIVE")
    title_run.font.color.rgb = RGBColor(36, 64, 97)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("System Architecture, Operations, and User Reference Guide")

    meta = add_table(
        document,
        ["Document Field", "Value"],
        [
            ["System Name", "KURA LIVE - Kenya Unified Results Architecture"],
            ["Primary Purpose", "Election monitoring, verification, live stream intelligence, and discrepancy management"],
            ["Deployment Context", "Next.js 14 local deployment with Supabase scaffolding and local fallback persistence"],
            ["Prepared For", "Engineers, election operators, field agents, administrators, and support teams"],
            ["Workspace", r"C:\Users\VincentOmbok\Documents\Kura27project"],
            ["Current Runtime", "Localhost deployment on port 3000"],
        ],
        [2.2, 4.3],
    )

    document.add_paragraph("")
    add_heading(document, "How To Use This Guide", 1)
    add_body(document, "This guide is written for two audiences at once. Engineers can use it to understand architecture, integration points, configuration, and maintenance. Non-technical users can use it to learn how the system works operationally and how to complete common tasks.")
    add_bullet(document, "If you need a quick orientation, start with System Overview, Features and Functionality, and User Guide.")
    add_bullet(document, "If you are deploying, troubleshooting, or extending the platform, read Architecture and Design, Technical Specifications, Installation and Setup, and Maintenance and Support.")
    add_bullet(document, "If you are monitoring live election operations, focus on the Agent Workflow, Command Centre, Comparison Engine, Alerts, and Troubleshooting sections.")

    add_heading(document, "Table Of Contents", 1)
    toc_items = [
        "1. System Overview",
        "2. Architecture and Design",
        "3. Technical Specifications",
        "4. Features and Functionality",
        "5. Installation and Setup",
        "6. Configuration Options",
        "7. User Guide",
        "8. API Reference and Integration Points",
        "9. Troubleshooting and FAQs",
        "10. Maintenance and Support",
    ]
    for item in toc_items:
        add_bullet(document, item)

    document.add_page_break()

    add_heading(document, "1. System Overview", 1)
    add_body(document, "KURA LIVE is an election monitoring and verification platform designed to observe live polling-station result announcements, convert spoken announcements into structured tallies, compare those tallies against scanned Form 34A results and official IEBC figures, and raise discrepancies in real time.")
    add_body(document, "The platform does not host or originate livestreams. Instead, field agents use their own TikTok accounts, YouTube accounts, or approved external tools to go live from polling stations. KURA LIVE then monitors those streams, captures shared tab audio, transcribes the result announcement workflow, and feeds the extracted figures into a three-source verification engine.")
    add_heading(document, "Core Purpose", 2)
    add_bullet(document, "Give field agents a structured way to register and monitor live polling-station broadcasts.")
    add_bullet(document, "Transcribe spoken vote announcements with OpenAI Whisper and parse them into candidate vote counts.")
    add_bullet(document, "Compare live audio tallies with OCR-extracted Form 34A tallies and official IEBC values.")
    add_bullet(document, "Flag any discrepancy of one vote or more and surface it nationally in a command centre.")
    add_bullet(document, "Support both live election operations and fallback demonstration/testing workflows on localhost.")
    add_heading(document, "Key Capabilities", 2)
    add_table(
        document,
        ["Capability", "What It Means In Practice"],
        [
            ["Field stream registration", "Agents paste TikTok, YouTube, or RTMP-related stream URLs into KURA LIVE and bind them to stations."],
            ["Live audio monitoring", "Operators use browser tab audio capture to ingest spoken announcements without illegally ripping source streams."],
            ["Realtime transcription", "Captured audio is chunked and sent to Whisper progressively, not only after a full clip ends."],
            ["Three-source verification", "Audio AI, Form 34A OCR, and IEBC official figures are stored and compared per candidate per station."],
            ["Alert escalation", "The system automatically creates open alerts when deltas are detected and supports escalation or dismissal workflows."],
            ["National command centre", "A geospatial and operational dashboard shows live streams, verification status, alerts, and tally movement nationally."],
        ],
        [2.15, 4.35],
    )

    add_heading(document, "2. Architecture and Design", 1)
    add_body(document, "KURA LIVE follows a single-repository Next.js application model. The frontend, API routes, UI state, and local fallback persistence live in the same codebase. Supabase is scaffolded as the long-term system of record, while a local JSON persistence layer is used to keep the demo and localhost environment stable even when hosted schema support is incomplete.")
    add_heading(document, "High-Level Component Model", 2)
    add_code_block(
        document,
        [
            "Field Agent Live Stream URL",
            "    -> Agent Portal (Stream Management + Transcription)",
            "    -> Browser Tab Audio Capture",
            "    -> Chunked Live Transcription API",
            "    -> Whisper Processing + Vote Parser",
            "    -> Tally Persistence",
            "    -> Reconciliation Engine",
            "    -> Alerts + Verification Events",
            "    -> Command Centre + Map + Stream Wall + Comparison Views",
        ],
    )
    add_heading(document, "Major Runtime Layers", 2)
    add_table(
        document,
        ["Layer", "Main Responsibilities", "Representative Files"],
        [
            ["Presentation layer", "Pages, route layouts, operator views, agent views, map view, stream wall, and public results UX.", "src/app/*, src/components/*"],
            ["Client state layer", "Realtime snapshot hydration, tallies, alerts, and toast notifications.", "src/store/websocketStore.ts, src/store/tallyStore.ts, src/store/alertStore.ts, src/store/toastStore.ts"],
            ["API layer", "Stream registration, transcription, forms processing, IEBC sync, alerts, tally comparison, and command-centre snapshots.", "src/app/api/*"],
            ["Domain logic layer", "Vote parsing, candidate matching, discrepancy scoring, live session control, command-centre snapshot building.", "src/lib/*.ts"],
            ["Persistence layer", "Supabase scaffolding for relational tables plus local JSON fallback state.", "src/lib/supabase.ts, src/lib/verificationDemo.ts, supabase/migrations/*"],
            ["Security and routing layer", "Session cookie parsing, path authorization, middleware redirects, and role separation.", "src/lib/session.ts, src/lib/access.ts, src/middleware.ts"],
        ],
        [1.25, 2.75, 2.5],
    )
    add_heading(document, "Primary User Experiences", 2)
    add_bullet(document, "Agent Portal: Field-facing workspace for stream monitoring, live transcription, Form 34A processing, comparison, and alerts.")
    add_bullet(document, "Admin Portal: Administrative and oversight workspace for broader campaign or operational controls.")
    add_bullet(document, "Command Centre: National monitoring dashboard showing stations, streams, alerts, and aggregation behavior.")
    add_bullet(document, "Public Results Portal: Public-facing view for results and map context.")

    add_heading(document, "3. Technical Specifications", 1)
    add_table(
        document,
        ["Category", "Specification"],
        [
            ["Frontend framework", "Next.js 14 with App Router"],
            ["Language", "TypeScript"],
            ["Styling", "Tailwind CSS"],
            ["UI state", "Zustand"],
            ["Forms and validation", "React Hook Form and Zod"],
            ["Animation", "Framer Motion"],
            ["Charting", "Recharts"],
            ["Video support", "react-player plus custom stream preview handling"],
            ["Map stack", "Leaflet and react-leaflet with SSR-safe dynamic imports"],
            ["Database scaffolding", "Supabase PostgreSQL"],
            ["AI transcription", "OpenAI Whisper API"],
            ["Deployment mode validated here", "Localhost deployment via Next.js production server"],
        ],
        [2.1, 4.4],
    )
    add_heading(document, "Important Code Modules", 2)
    add_table(
        document,
        ["Module", "Role In The System"],
        [
            ["src/components/agent/TranscriptionPanel.tsx", "Handles demo audio, live tab capture, chunk uploads, transcript display, and extracted vote updates."],
            ["src/components/agent/StreamPanel.tsx", "Manages stream URL inputs, platform preference, preview, and live stream context for agents."],
            ["src/components/command/LiveMapPanel.tsx", "National map view with station state, verification markers, and stream/alert overlays."],
            ["src/components/command/StreamWall.tsx", "Displays active live sessions with operational metrics and quick feed access."],
            ["src/lib/verificationDemo.ts", "Local fallback persistence, seeded demo state, alert synchronization, and comparison-row generation."],
            ["src/lib/commandCenter.ts", "Builds the normalized command-centre snapshot consumed by the national dashboard."],
            ["src/app/api/transcription/live/route.ts", "Processes 5 to 10 second live audio chunks and returns progressive transcript/tally updates."],
            ["src/app/api/streams/route.ts", "Registers streams, stores or mirrors live sessions, and updates lifecycle state."],
        ],
        [2.0, 4.5],
    )

    add_heading(document, "4. Features and Functionality", 1)
    add_heading(document, "4.1 Stream Management", 2)
    add_body(document, "Agents can supply TikTok Live URLs, YouTube Live URLs, or platform URLs associated with their polling station. The system validates and stores the stream association, then opens a monitoring preview without pretending to own the underlying stream provider.")
    add_bullet(document, "Primary operator action: paste a live URL and choose the platform context.")
    add_bullet(document, "System action: register the stream session, bind it to agent identity and station code, and expose it to the command centre.")
    add_bullet(document, "Operational limitation: the system uses browser-shared audio capture instead of extracting media directly from restricted platforms.")

    add_heading(document, "4.2 Live Audio Capture and Transcription", 2)
    add_body(document, "The agent starts live monitoring from the Transcription panel. The browser asks for tab or system audio access. Once approved, KURA LIVE uses the MediaRecorder API to create progressive audio chunks, normally on an 8-second cadence in the current implementation. Each chunk is uploaded to the live transcription route and processed independently.")
    add_table(
        document,
        ["Step", "Current Behavior"],
        [
            ["Capture permission", "User shares a browser tab with audio enabled."],
            ["Chunk creation", "MediaRecorder emits a new audio blob approximately every 8 seconds."],
            ["API handling", "The chunk is sent to /api/transcription/live with session, station, and timing metadata."],
            ["AI processing", "OpenAI Whisper is used when configured; otherwise or on failure, the system falls back gracefully to deterministic parsing behavior."],
            ["UI update", "Transcript text, extracted votes, tallies, and alerts update progressively in the agent portal."],
        ],
        [1.6, 4.9],
    )

    add_heading(document, "4.3 Vote Extraction and Parsing", 2)
    add_body(document, "The parsing layer converts freeform announcements into structured results. It supports English and partial Swahili phrases, resolves candidate aliases, suppresses duplicates, and assigns confidence scores. Rejected votes are also detected when phrased explicitly.")
    add_bullet(document, "Example English pattern: John Kariuki received 120 votes.")
    add_bullet(document, "Example Swahili pattern: Mary Wambui amepata kura 115.")
    add_bullet(document, "Edge behavior: if Whisper fails, the system can still parse fallback transcript text and keep the operator workflow moving.")

    add_heading(document, "4.4 Form 34A OCR Flow", 2)
    add_body(document, "The Form 34A module accepts bundled samples or user uploads and maps them into OCR-style extracted results for the demonstration environment. These results are stored as the Form 34A verification source and immediately enter the comparison matrix.")

    add_heading(document, "4.5 IEBC Official Source Integration", 2)
    add_body(document, "The current implementation uses a mock official dataset to simulate the IEBC portal feed. The API persists those values as the IEBC source so that three-way comparison can happen deterministically in the local environment.")

    add_heading(document, "4.6 Comparison and Alerts", 2)
    add_body(document, "For every candidate at a station, the system compares Audio AI, Form 34A, and IEBC tallies. Any discrepancy of one vote or more creates an alert. Severity is warning for deltas from one to four and critical for five or more.")
    add_bullet(document, "Open alerts appear in the agent portal and command centre.")
    add_bullet(document, "Operators can escalate or dismiss alerts from the alerts panel.")
    add_bullet(document, "Verification events and station activity are also tracked in local state and command-centre snapshots.")

    add_heading(document, "4.7 Command Centre and Map View", 2)
    add_body(document, "The command centre at /dashboard is the national monitoring plane. It combines a live station map, tally summary, incident panel, stream wall, AI monitoring panel, and an operational right rail. The map view at /map-view gives a larger, map-focused monitoring screen.")
    add_table(
        document,
        ["Dashboard Area", "Operational Use"],
        [
            ["Live map", "Shows station status, stream presence, verification health, and alert counts geographically."],
            ["Stream wall", "Lists monitored streams with platform, signal quality, latency, and operator access controls."],
            ["Results tally", "Aggregates candidate performance and turnout indicators for top-level monitoring."],
            ["Incident and alert panels", "Surface discrepancies and operational problems needing review."],
            ["Right rail", "Displays high-level system health, turnout, counties reporting, and API/stream health indicators."],
        ],
        [2.0, 4.5],
    )

    add_heading(document, "5. Installation and Setup", 1)
    add_number(document, "Install dependencies with npm install.")
    add_number(document, "Create or update .env.local with the required environment variables.")
    add_number(document, "Run npm run build to generate a production build.")
    add_number(document, "Run npm run start -- --port 3000 for the local production server, or npm run dev for development.")
    add_number(document, "Open http://localhost:3000 in a browser.")
    add_heading(document, "Required Local Commands", 2)
    add_code_block(
        document,
        [
            "npm install",
            "npm run lint",
            "npm run typecheck",
            "npm test",
            "npm run build",
            "npm run start -- --port 3000",
        ],
    )
    add_heading(document, "Supabase Preparation", 2)
    add_body(document, "Two migration files are present for the verification and live monitoring model. These should be applied to the Supabase project if the hosted database is intended to replace local fallback persistence.")
    add_bullet(document, "supabase/migrations/20260504_verification_pipeline.sql")
    add_bullet(document, "supabase/migrations/20260511_live_monitoring_pipeline.sql")

    add_heading(document, "6. Configuration Options", 1)
    add_table(
        document,
        ["Variable", "Purpose", "Required"],
        [
            ["NEXT_PUBLIC_SUPABASE_URL", "Supabase project URL used by browser and server-side client scaffolding.", "Yes"],
            ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Public Supabase anon key for client-side access.", "Yes"],
            ["SUPABASE_SERVICE_ROLE_KEY", "Server-side elevated Supabase key for privileged operations.", "Yes for full DB integration"],
            ["NEXT_PUBLIC_APP_URL", "Base application URL used for local routing and references.", "Yes"],
            ["OPENAI_API_KEY", "Used for Whisper transcription of uploaded and live chunked audio.", "Yes for real Whisper use"],
        ],
        [2.05, 3.75, 0.9],
    )
    add_body(document, "Important security rule: service keys and OpenAI keys must never be exposed in browser code, screenshots, or public documentation. Only server routes should use them.")

    add_heading(document, "7. User Guide", 1)
    add_heading(document, "7.1 Field Agent Workflow", 2)
    add_number(document, "Sign in to the agent portal.")
    add_number(document, "Open the Live Stream tab and register or preview the monitoring URL.")
    add_number(document, "Open the Transcription tab.")
    add_number(document, "Use Open agent TikTok, Open agent YouTube, or paste a live URL.")
    add_number(document, "Click Start live capture and share the correct browser tab with audio enabled.")
    add_number(document, "Allow at least one chunk cycle so transcript and tally updates appear.")
    add_number(document, "Upload or process a Form 34A sample from the Form 34A tab.")
    add_number(document, "Fetch IEBC tallies.")
    add_number(document, "Review Comparison and Alerts.")
    add_heading(document, "7.2 Command Centre Operator Workflow", 2)
    add_number(document, "Open /dashboard after authentication.")
    add_number(document, "Review the map for reporting stations, discrepancy hotspots, and stream density.")
    add_number(document, "Open streams from the stream wall if a station needs deeper review.")
    add_number(document, "Use the tally panel for a quick national summary.")
    add_number(document, "Review incidents and alerts for escalation priority.")
    add_number(document, "Use the full map at /map-view for focused geospatial monitoring.")
    add_heading(document, "7.3 Public Viewer Workflow", 2)
    add_number(document, "Open /results to view public-facing results and map context.")
    add_number(document, "Use search and filtering to inspect station or constituency context.")
    add_number(document, "Use /results-tally for a denser tally-specific presentation.")

    add_heading(document, "8. API Reference and Integration Points", 1)
    add_table(
        document,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/api/streams", "GET / POST / PATCH", "List streams, register new monitoring sessions, and update lifecycle state."],
            ["/api/transcription", "POST", "Process uploaded or demo audio as a non-live transcription flow."],
            ["/api/transcription/live", "POST", "Process chunked live audio with session metadata and return progressive transcript updates."],
            ["/api/forms", "GET / POST", "List form uploads and process Form 34A samples or uploads."],
            ["/api/iebc/[station]", "GET", "Fetch and optionally persist official IEBC comparison data."],
            ["/api/tallies", "GET / POST", "Read tally entries or persist source-specific values."],
            ["/api/alerts", "GET / PATCH", "Read alerts and update alert status."],
            ["/api/verification/compare", "GET / POST", "Read comparison rows and station verification stats."],
            ["/api/command-centre", "GET", "Return the normalized national dashboard snapshot."],
        ],
        [2.25, 1.25, 3.0],
    )
    add_heading(document, "Typical Live Transcription Payload", 2)
    add_code_block(
        document,
        [
            "station_code=KE-047-290-0001",
            "agent_id=agent-001",
            "session_id=<uuid>",
            "chunk_index=0",
            "started_at=<unix ms>",
            "ended_at=<unix ms>",
            "sample_id=audio-westlands-official",
            "audio=<binary webm blob>",
        ],
    )

    add_heading(document, "9. Troubleshooting and FAQs", 1)
    add_heading(document, "Common Issues", 2)
    add_table(
        document,
        ["Problem", "Likely Cause", "Recommended Action"],
        [
            ["Command centre returns an error", "Persisted local demo state uses an older schema.", "Restart the app after the state loader self-heals, or reset the local demo bootstrap endpoint."],
            ["No live transcript appears", "Tab audio was not shared, browser denied capture, or the source tab has no audible content.", "Restart live capture and ensure the selected tab is shared with audio enabled."],
            ["Whisper transcription does not use the real API", "OPENAI_API_KEY is missing or the upstream request failed.", "Verify the environment variable and check network/API availability."],
            ["Form processing works but does not reflect a real OCR engine", "The localhost environment uses sample-based extraction fallback.", "Apply Supabase and OCR backend integrations if production OCR is required."],
            ["Dashboard is accessible in raw HTTP tests but not in the browser session", "The browser session may not be authenticated or cookies may have expired.", "Log in again and confirm the correct role cookie is present."],
            ["Build or typecheck runs out of memory", "Node heap was too low for this workspace.", "Use the package scripts provided; they already raise the heap limit to 4096 MB."],
        ],
        [1.9, 2.0, 2.6],
    )
    add_heading(document, "Frequently Asked Questions", 2)
    add_bullet(document, "Does KURA LIVE host the stream? No. KURA LIVE monitors streams owned and started elsewhere.")
    add_bullet(document, "Can the system directly rip TikTok media? No. The supported compliance model is browser-shared audio capture.")
    add_bullet(document, "Can agents see the command centre? Yes, the shared command centre is accessible to authenticated agents for monitoring.")
    add_bullet(document, "Is Supabase fully authoritative today? Not in this localhost rescue state. The code supports Supabase, but local JSON fallback persistence remains active when hosted schema support is incomplete.")
    add_bullet(document, "Can this guide be exported to PDF? Yes. The generated Word document is designed for Word or PDF export.")

    add_heading(document, "10. Maintenance and Support", 1)
    add_heading(document, "Routine Maintenance", 2)
    add_bullet(document, "Re-run lint, typecheck, tests, and build after any code change.")
    add_bullet(document, "Apply pending Supabase migrations if the hosted database is being used as the authoritative store.")
    add_bullet(document, "Rotate secrets such as OpenAI keys and service-role keys according to your security policy.")
    add_bullet(document, "Review the command-centre snapshot output when adding new live-session or verification fields.")
    add_bullet(document, "Keep browser-capture flows tested in the target browser versions used by field agents.")
    add_heading(document, "Operational Escalation Path", 2)
    add_number(document, "Field agent problem: confirm login, station assignment, and tab-audio permissions.")
    add_number(document, "Application problem: check API route responses, alert/tally state, and local persistence health.")
    add_number(document, "Infrastructure problem: check environment variables, Supabase connectivity, and Whisper API availability.")
    add_number(document, "Production support problem: promote issue details to the engineering team with route, station, session, and transcript timing context.")
    add_heading(document, "Validation Baseline", 2)
    add_body(document, "At the time of this documentation build, the validated engineering checks are npm run lint, npm run typecheck, npm test, and npm run build. The local production server is expected to run on port 3000.")

    document.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(document, "Appendix A: Key Local Routes", 1)
    add_table(
        document,
        ["Route", "Audience", "Purpose"],
        [
            ["http://localhost:3000/", "All users", "Landing page and navigation starting point."],
            ["http://localhost:3000/login", "Authenticated users", "Main login route for shared credentials."],
            ["http://localhost:3000/agent-login", "Field agents", "Dedicated agent login flow."],
            ["http://localhost:3000/agent", "Field agents", "Agent portal for live stream monitoring and verification."],
            ["http://localhost:3000/dashboard", "Agents and admins", "Shared command centre."],
            ["http://localhost:3000/admin", "Admins only", "Administrative oversight portal."],
            ["http://localhost:3000/map-view", "Agents and admins", "Full-page national map monitoring view."],
            ["http://localhost:3000/results", "Public and operators", "Public results and contextual monitoring view."],
        ],
        [2.25, 1.3, 3.2],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT))


if __name__ == "__main__":
    build_document()
